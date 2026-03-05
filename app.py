import streamlit as st
import os
import sys
import re
import requests
from io import BytesIO 
from docx import Document 
from dotenv import load_dotenv
from supabase import create_client
from pypdf import PdfReader

# Dem Server den Weg zum 'backend' Ordner zeigen
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from backend.shared.utils.docx_builder import build_cover_letter_docx, SwissLetterHeader
from backend.api.llm.prompt_builder import build_single_call_prompt
from backend.api.llm.llm_client import generate_cover_letter_body_only

# ==========================================
# 1. INITIALISIERUNG & SETUP
# ==========================================
load_dotenv()
st.set_page_config(page_title="Swiss-Apply KI", page_icon="🇨🇭", layout="wide")

# ==========================================
# 1.5 STARTUP-DESIGN (CSS-Tricks)
# ==========================================
hide_st_style = """
<style>
/* Versteckt den Deploy-Button und das Menü oben rechts */
[data-testid="stToolbar"] {visibility: hidden !important;}
/* Versteckt den Streamlit-Footer ganz unten */
footer {visibility: hidden !important;}
/* Macht den weissen Abstand ganz oben etwas schmaler */
.block-container {padding-top: 2rem;}
</style>
"""
st.markdown(hide_st_style, unsafe_allow_html=True)

# Verbindung zu Supabase Zürich
supabase_url = os.getenv("SUPABASE_URL")
supabase_key = os.getenv("SUPABASE_KEY")
supabase = create_client(supabase_url, supabase_key)

# ==========================================
# 2. STREAMLIT "GEDÄCHTNIS" (SESSION STATE)
# ==========================================
if 'user' not in st.session_state:
    st.session_state.user = None

# Supabase den Login-Token bei jedem Klick wieder übergeben
if 'access_token' in st.session_state and 'refresh_token' in st.session_state:
    try:
        supabase.auth.set_session(st.session_state.access_token, st.session_state.refresh_token)
    except Exception:
        pass

# ==========================================
# 2.5 PASSWORT-RESET ABFANGEN
# ==========================================
if "code" in st.query_params:
    try:
        res = supabase.auth.exchange_code_for_session({"auth_code": st.query_params["code"]})
        st.session_state.user = res.user
        st.session_state.access_token = res.session.access_token
        st.session_state.refresh_token = res.session.refresh_token
        st.query_params.clear()
        st.success("✅ Verifizierung erfolgreich! Bitte setze jetzt unten ein neues Passwort.")
    except Exception:
        st.error("Der Link ist ungültig oder abgelaufen.")

# ==========================================
# 3. HILFSFUNKTIONEN
# ==========================================
def create_rav_pdf_lookalike(data):
    doc = Document()
    doc.add_heading('Nachweis der persönlichen Arbeitsbemühungen', 0)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Datum'
    hdr_cells[1].text = 'Firma, Ort'
    hdr_cells[2].text = 'Art der Stelle'
    hdr_cells[3].text = 'Form'
    hdr_cells[4].text = 'Ergebnis'
    for item in data:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item.get('applied_at', ''))[:10]
        row_cells[1].text = item.get('company_name', '')
        row_cells[2].text = item.get('job_title', '')
        row_cells[3].text = item.get('channel', 'Online')
        row_cells[4].text = "In Prüfung"
    file_stream = BytesIO()
    doc.save(file_stream)
    return file_stream.getvalue()

# ==========================================
# 4. AUTHENTIFIZIERUNG (LOGIN / REGISTRIERUNG)
# ==========================================
if st.session_state.user is None:
    st.title("🇨🇭 Swiss-Apply: Deine Bewerbungs-Plattform")
    t_login, t_reg, t_pw = st.tabs(["🔐 Login", "📝 Registrieren", "🔑 Passwort vergessen"])

    with t_login:
        with st.form("login_form"):
            email_in = st.text_input("E-Mail")
            pass_in = st.text_input("Passwort", type="password")
            if st.form_submit_button("Anmelden", use_container_width=True):
                try:
                    res = supabase.auth.sign_in_with_password({"email": email_in, "password": pass_in})
                    st.session_state.user = res.user
                    st.session_state.access_token = res.session.access_token
                    st.session_state.refresh_token = res.session.refresh_token
                    st.rerun()
                except Exception as e:
                    st.error(f"Login fehlgeschlagen. Grund: {e}")

    with t_reg:
        with st.form("reg_form"):
            new_email = st.text_input("Deine E-Mail")
            new_pass = st.text_input("Passwort (min. 6 Zeichen)", type="password")
            invite_code = st.text_input("Geheimer Einladungscode", type="password", help="Frag Carlos nach dem Code!")
            
            if st.form_submit_button("Konto erstellen", use_container_width=True):
                if invite_code != "SWISS2026":
                    st.error("❌ Falscher Einladungscode! Registrierung blockiert.")
                elif len(new_pass) < 6:
                    st.error("Das Passwort muss mindestens 6 Zeichen lang sein.")
                else:
                    try:
                        supabase.auth.sign_up({"email": new_email, "password": new_pass})
                        st.success("✅ Registrierung erfolgreich! (Falls 'Confirm Email' aktiv ist, prüfe dein Postfach)")
                    except Exception as e: 
                        st.error(f"Fehler: {e}")

    with t_pw:
        with st.form("pw_form"):
            reset_email = st.text_input("E-Mail für Reset-Link")
            if st.form_submit_button("Link senden", use_container_width=True):
                try:
                    supabase.auth.reset_password_for_email(reset_email)
                    st.info("Check dein Postfach!")
                except Exception as e: 
                    st.error(f"Fehler beim Senden: {e}")

# ==========================================
# 5. HAUPTPROGRAMM (EINGELOGGT)
# ==========================================
else:
    st.sidebar.title("🇨🇭 Swiss-Apply")
    st.sidebar.write(f"Eingeloggt: **{st.session_state.user.email}**")
    if st.sidebar.button("Logout"):
        supabase.auth.sign_out()
        st.session_state.user = None
        st.session_state.pop('access_token', None)
        st.session_state.pop('refresh_token', None)
        st.rerun()

    # --- PROFIL-DATEN SICHER LADEN ---
    try:
        profile_res = supabase.table("profiles").select("*").eq("id", st.session_state.user.id).single().execute()
        p_data = profile_res.data if profile_res.data else {}
    except Exception:
        p_data = {}

    current_name = p_data.get("full_name", "")
    current_street = p_data.get("street", "")
    current_city = p_data.get("city", "")
    current_phone = p_data.get("phone", "")
    current_resume = p_data.get("resume_text", "")

    # --- PROFIL & LEBENSLAUF UPLOAD ---
    with st.expander("👤 Mein Profil & Lebenslauf"):
        with st.form("p_form"):
            st.subheader("1. Kontaktdaten")
            c1, c2 = st.columns(2)
            f_name = c1.text_input("Vollständiger Name", value=current_name)
            f_street = c1.text_input("Strasse & Nr.", value=current_street)
            f_city = c2.text_input("PLZ & Ort", value=current_city)
            f_phone = c2.text_input("Telefon", value=current_phone)
            
            st.subheader("2. Lebenslauf (Für die KI)")
            if current_resume:
                st.success("✅ Ein Lebenslauf ist bereits gespeichert! (Du kannst ihn überschreiben, indem du einen neuen hochlädst)")
            
            uploaded_pdf = st.file_uploader("Lebenslauf hochladen (Nur PDF)", type=["pdf"])
            
            if st.form_submit_button("Profil & Lebenslauf speichern"):
                try:
                    update_data = {
                        "id": st.session_state.user.id, 
                        "full_name": f_name, 
                        "street": f_street, 
                        "city": f_city, 
                        "phone": f_phone
                    }
                    
                    if uploaded_pdf is not None:
                        reader = PdfReader(uploaded_pdf)
                        extracted_text = ""
                        for page in reader.pages:
                            extracted_text += page.extract_text() + "\n"
                        update_data["resume_text"] = extracted_text
                        st.info("PDF erfolgreich gelesen und Text extrahiert!")

                    supabase.table("profiles").upsert(update_data).execute()
                    st.success("Profil erfolgreich aktualisiert!")
                    st.rerun()
                except Exception as e:
                    st.error(f"Fehler beim Speichern: {e}")

    # --- PASSWORT ÄNDERN ---
    with st.expander("🔑 Mein Passwort ändern"):
        with st.form("pwd_reset_form"):
            new_pwd = st.text_input("Neues Passwort (min. 6 Zeichen)", type="password")
            if st.form_submit_button("Passwort aktualisieren"):
                try:
                    supabase.auth.update_user({"password": new_pwd})
                    st.success("Passwort erfolgreich geändert! Du kannst dich ab sofort damit einloggen.")
                except Exception as e:
                    st.error(f"Fehler beim Ändern: {e}")

   # --- LIVE JOB-BOARD (API) ---
    st.divider()
    st.header("🎯 Live Job-Börse")
    st.write("Finde aktuelle Stelleninserate direkt hier in der App.")
    
    cj1, cj2 = st.columns(2)
    api_job = cj1.text_input("Welchen Job suchst du?", placeholder="z.B. Projektleiter", key="api_job")
    api_city = cj2.text_input("Wo?", value=current_city if current_city else "Schweiz", key="api_city")
    
    # Dein Key ist hier korrekt hinterlegt
    JOOBLE_API_KEY = "19414f40-206a-4cf6-ac1b-0106e96f7f8d"
    
    if st.button("Jobs live laden"):
        # Hier prüfen wir jetzt nur noch, ob der Platzhalter-Text weg ist
        if JOOBLE_API_KEY == "" or "DEIN_JOOBLE_KEY" in JOOBLE_API_KEY:
            st.warning("⚠️ Bitte trage deinen Jooble-API-Schlüssel im Code ein!")
        else:
            with st.spinner("Verbinde mit Job-Servern..."):
                try:
                    url = f"https://ch.jooble.org/api/{JOOBLE_API_KEY}"
                    payload = {"keywords": api_job, "location": api_city}
                    headers = {"Content-type": "application/json"}
                    response = requests.post(url, json=payload, headers=headers)
                    
                    if response.status_code == 200:
                        jobs = response.json().get("jobs", [])
                        if not jobs:
                            st.info("Für diese Suchanfrage wurden leider keine Jobs gefunden.")
                        else:
                            st.success(f"✅ {len(jobs)} brandneue Stellen gefunden!")
                            for job in jobs[:5]:
                                with st.container():
                                    st.subheader(job.get("title", "Kein Titel"))
                                    st.write(f"🏢 **{job.get('company', 'Firma unbekannt')}** | 📍 {job.get('location', '')}")
                                    snippet = job.get("snippet", "").replace("<b>", "**").replace("</b>", "**")
                                    st.write(f"> {snippet}")
                                    st.link_button("Komplettes Inserat ansehen", job.get("link", "#"))
                                    st.divider()
                    else:
                        st.error(f"Fehler vom Server. Code: {response.status_code}. (Hinweis: Prüfe, ob dein Jooble-Key bereits per E-Mail bestätigt wurde.)")
                except Exception as e:
                    st.error(f"Technischer Fehler beim Abruf: {e}")

    # --- GENERATOR ---
    st.divider()
    st.header("🚀 Neuer Bewerbungs-Brief")
    st.write("Kopiere die Daten deines Wunsch-Jobs hier hinein.")
    col_a, col_b = st.columns(2)
    company = col_a.text_input("Firma", placeholder="z.B. Swisscom")
    job_title = col_b.text_input("Stelle", placeholder="z.B. Projektleiter")
    job_desc = st.text_area("Stellenbeschreibung (Kopiere den Text aus dem Inserat)", height=150)

    if st.button("Brief generieren & loggen"):
        if not company or not job_title:
            st.warning("Firma und Stelle fehlen!")
        else:
            with st.spinner("KI arbeitet..."):
                try:
                    header = SwissLetterHeader(
                        full_name=current_name if current_name else "Name fehlt",
                        street=current_street if current_street else "Strasse fehlt",
                        postal_code="", city=current_city if current_city else "Ort fehlt",
                        email=st.session_state.user.email
                    )
                    
                    prompt = build_single_call_prompt(
                        {"personal": {"full_name": current_name}}, 
                        {"title": job_title, "company": company, "location": "CH", "requirements": {"keywords": [job_desc[:50]]}}
                    )
                    
                    if current_resume:
                        prompt += f"\n\nWICHTIGE ZUSATZINFO: Hier ist der Lebenslauf des Bewerbers. Bitte beziehe unbedingt echte Erfahrungen, Arbeitgeber und Fähigkeiten aus diesem Lebenslauf in das Anschreiben ein, sofern sie zur Stelle passen. Erfinde keine Erfahrungen, die nicht im Lebenslauf stehen!\n\nLEBENSLAUF:\n{current_resume}"
                    
                    llm_result = generate_cover_letter_body_only(prompt=prompt)
                    
                    docx_bytes = build_cover_letter_docx(
                        header=header, recipient_block=f"{company}\nPersonalabteilung",
                        place_and_date=f"{current_city or 'Schweiz'}, 05.03.2026",
                        subject=f"Bewerbung als {job_title}",
                        body_text=llm_result.body_text, signature_name=current_name
                    )
                    
                    # Cloud Upload
                    clean_name = re.sub(r'[^a-zA-Z0-9]', '_', company)
                    path = f"web_generiert/{st.session_state.user.id}/{clean_name}.docx"
                    supabase.storage.from_("letters").upload(path=path, file=docx_bytes, file_options={"upsert": "true"})

                    # In DB loggen
                    supabase.table("applications").insert({
                        "user_id": st.session_state.user.id, "company_name": company,
                        "job_title": job_title, "status": "generated", "channel": "Online"
                    }).execute()
                    
                    st.success("Brief generiert und sicher archiviert!")
                    st.download_button("📂 Download Word", data=docx_bytes, file_name=f"Bewerbung_{clean_name}.docx")
                except Exception as e:
                    st.error(f"Fehler bei der Generierung oder Speicherung: {e}")

    # --- BEWERBUNGS-HISTORIE & RAV DASHBOARD ---
    st.divider()
    st.header("📊 Dein RAV-Dashboard & Historie")
    
    try:
        apps_res = supabase.table("applications").select("*").eq("user_id", st.session_state.user.id).order("applied_at", desc=True).execute()
        apps = apps_res.data if apps_res.data else []
        
        c_met1, c_met2 = st.columns(2)
        c_met1.metric("Bewerbungen (Gesamt)", len(apps), f"{len(apps)}/10 Ziel")
        c_met2.progress(min(len(apps)/10, 1.0), text="Fortschritt RAV-Ziel")

        if apps:
            if st.button("📄 Offizielles RAV-Beiblatt erstellen"):
                rav_docx = create_rav_pdf_lookalike(apps)
                st.download_button("📥 RAV-Nachweis herunterladen", data=rav_docx, file_name="RAV_Nachweis.docx")
    except Exception as e:
        st.warning(f"RAV-Daten konnten nicht geladen werden: {e}")

    st.subheader("📂 Deine Dokumenten-Ablage")
    try:
        user_id_path = f"web_generiert/{st.session_state.user.id}/"
        files = supabase.storage.from_("letters").list(path=user_id_path)
        
        if files:
            for f in files:
                if f['name'] == '.empty': continue
                col_h1, col_h2 = st.columns([3, 1])
                with col_h1:
                    st.write(f"📄 **{f['name']}**")
                with col_h2:
                    file_data = supabase.storage.from_("letters").download(f"{user_id_path}{f['name']}")
                    st.download_button("📥 Laden", data=file_data, file_name=f['name'], key=f"dl_{f['name']}")
        else:
            st.info("Noch keine Dokumente archiviert.")
    except Exception:
        st.write("Erstelle deine erste Bewerbung, um die Ablage zu aktivieren.")

    with st.expander("🛡️ Datenschutz-Info"):
        st.write("Deine Daten liegen verschlüsselt in der Region **Zürich (eu-central-2)**.")