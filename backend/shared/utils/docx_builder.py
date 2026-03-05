from __future__ import annotations
from dataclasses import dataclass
from typing import Optional
from io import BytesIO
import zipfile
import datetime

import docx
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.shared import OxmlElement, qn

"""
DOCX Builder (Swiss format) — SN 010130 Layout + Forensische Tarnkappen-Funktion
Verschmelzung: ChatGPT Cloud-Architecture + Swiss Typography
"""

@dataclass
class SwissLetterHeader:
    full_name: str
    street: str
    postal_code: str
    city: str
    phone: Optional[str] = None
    email: Optional[str] = None

def add_hyperlink(paragraph, text, url, font_name="Arial", font_size=11):
    """Fügt einen klickbaren, blauen Link ein und erzwingt die Schriftart."""
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.append(rFonts)
    
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(font_size * 2)))
    rPr.append(sz)
    
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0000FF')
    rPr.append(c)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return paragraph

def build_cover_letter_docx(
    header: SwissLetterHeader,
    recipient_block: str,
    place_and_date: str,
    subject: str,
    body_text: str,
    font_name: str = "Arial",
    closing: str = "Freundliche Grüsse",
    signature_name: Optional[str] = None,
) -> bytes:
    doc = Document()

    # --- SN 010130 Seitenränder (3.0cm links für CH-Lochung!) ---
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)

    # --- Dynamische Typografie ---
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15

    # --- ABSENDER BLOCK ---
    p_name = doc.add_paragraph()
    run_name = p_name.add_run(header.full_name)
    run_name.font.size = Pt(15)
    run_name.font.name = font_name
    run_name.bold = True
    p_name.paragraph_format.space_after = Pt(0) 
    
    p_abs = doc.add_paragraph()
    p_abs.paragraph_format.space_after = Pt(40) # Exakter Abstand zum Fenster
    
    # Adresse und Telefon
    abs_text = f"{header.street}\n{header.postal_code} {header.city}"
    if header.phone:
        abs_text += f"\nNatel {header.phone}"
    if header.email:
        abs_text += f"\nE-Mail "
        
    run_abs = p_abs.add_run(abs_text)
    run_abs.font.size = Pt(11)
    run_abs.font.name = font_name
    
    # E-Mail als Hyperlink
    if header.email:
        add_hyperlink(p_abs, header.email, f"mailto:{header.email}", font_name=font_name, font_size=11)

    # --- EMPFÄNGER BLOCK ---
    p_empf = doc.add_paragraph()
    p_empf.paragraph_format.space_after = Pt(28) 
    run_empf = p_empf.add_run(recipient_block)
    run_empf.font.size = Pt(11)
    run_empf.font.name = font_name

    # --- DATUM (Modern CH: Linksbündig) ---
    p_date = doc.add_paragraph(place_and_date)
    p_date.paragraph_format.space_after = Pt(24) 

    # --- BETREFF ---
    p_betr = doc.add_paragraph()
    p_betr.paragraph_format.space_after = Pt(24) 
    run_b = p_betr.add_run(subject)
    run_b.bold = True
    run_b.font.size = Pt(12)
    run_b.font.name = font_name

    # --- ANREDE ---
    p_anrede = doc.add_paragraph("Sehr geehrte Damen und Herren") 
    p_anrede.paragraph_format.space_after = Pt(12) 

    # --- TEXTKÖRPER ---
    for para in [p.strip() for p in body_text.split("\n\n") if p.strip()]:
        p_text = doc.add_paragraph(para)
        p_text.paragraph_format.space_after = Pt(12) 

    # --- ABSCHLUSS ---
    p_gruss = doc.add_paragraph(closing)
    p_gruss.paragraph_format.space_before = Pt(12)
    p_gruss.paragraph_format.space_after = Pt(36) 
    
    doc.add_paragraph(signature_name or header.full_name)

    # In Memory speichern
    bio = BytesIO()
    doc.save(bio)
    
    # Metadaten forensisch reinigen und als Bytes für Supabase zurückgeben
    return sanitize_docx_metadata(bio.getvalue())


# ==========================================
# FORENSISCHE METADATEN-REINIGUNG (ChatGPTs Meisterstück)
# ==========================================
def sanitize_docx_metadata(docx_bytes: bytes) -> bytes:
    core_xml = _minimal_core_xml()
    app_xml = _minimal_app_xml()

    in_mem = BytesIO(docx_bytes)
    out_mem = BytesIO()

    with zipfile.ZipFile(in_mem, "r") as zin, zipfile.ZipFile(out_mem, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            name = item.filename
            if name == "docProps/custom.xml":
                continue
            if name == "docProps/core.xml":
                zout.writestr(name, core_xml); continue
            if name == "docProps/app.xml":
                zout.writestr(name, app_xml); continue
            zout.writestr(name, zin.read(name))
    return out_mem.getvalue()

def _minimal_core_xml() -> str:
    now = datetime.datetime.utcnow().replace(microsecond=0).isoformat() + "Z"
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
 xmlns:dc="http://purl.org/dc/elements/1.1/"
 xmlns:dcterms="http://purl.org/dc/terms/"
 xmlns:dcmitype="http://purl.org/dc/dcmitype/"
 xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">
  <dcterms:created xsi:type="dcterms:W3CDTF">{now}</dcterms:created>
  <dcterms:modified xsi:type="dcterms:W3CDTF">{now}</dcterms:modified>
</cp:coreProperties>
"""

def _minimal_app_xml() -> str:
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"
 xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">
  <Application>Microsoft Office Word</Application>
  <DocSecurity>0</DocSecurity>
  <ScaleCrop>false</ScaleCrop>
  <LinksUpToDate>false</LinksUpToDate>
  <SharedDoc>false</SharedDoc>
  <HyperlinksChanged>false</HyperlinksChanged>
  <AppVersion>16.0000</AppVersion>
</Properties>
"""